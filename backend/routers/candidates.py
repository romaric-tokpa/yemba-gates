"""
Routes pour la gestion des candidats (US04)
"""
import os
import shutil
import json
import tempfile
import base64
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlmodel import Session, select
from sqlalchemy import text, func
from typing import List, Optional
from uuid import UUID, uuid4

# Imports pour l'extraction de texte
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

# Import pour Google Gemini
try:
    import google.generativeai as genai
    GeminiError = Exception
    try:
        # Tentative d'import des erreurs spécifiques de Gemini
        from google.api_core import exceptions as google_exceptions
        GeminiError = google_exceptions.GoogleAPIError
    except ImportError:
        pass
except ImportError:
    genai = None
    GeminiError = Exception

from database_tenant import get_session
from models import Candidate, User, UserRole, Interview, Application, Job, CandidateJobComparison
from schemas import CandidateCreate, CandidateUpdate, CandidateResponse, CandidateParseResponse, JobCandidateComparisonResponse
from auth import get_current_active_user, require_recruteur, require_client

router = APIRouter(prefix="/candidates", tags=["candidates"])

# Dossiers pour stocker les fichiers
UPLOAD_DIR = Path("uploads/cvs")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Dossier pour les photos de profil (dans static/uploads)
PHOTOS_DIR = Path("static/uploads")
PHOTOS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx"}
ALLOWED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp"}


def is_allowed_file(filename: str) -> bool:
    """Vérifie si le fichier a une extension autorisée"""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def is_allowed_image(filename: str) -> bool:
    """Vérifie si le fichier est une image autorisée"""
    return Path(filename).suffix.lower() in ALLOWED_IMAGE_EXTENSIONS


def get_creator_info(creator_id: UUID, session: Session) -> dict:
    """Récupère les informations du créateur (nom, prénom, email)"""
    creator = session.get(User, creator_id)
    if creator:
        return {
            "creator_first_name": creator.first_name,
            "creator_last_name": creator.last_name,
            "creator_email": creator.email,
        }
    return {
        "creator_first_name": None,
        "creator_last_name": None,
        "creator_email": None,
    }


def check_duplicate_candidate(
    session: Session,
    email: Optional[str] = None,
    first_name: Optional[str] = None,
    last_name: Optional[str] = None,
    phone: Optional[str] = None
) -> Optional[Candidate]:
    """
    Vérifie si un candidat avec les mêmes informations existe déjà
    
    Priorité de vérification:
    1. Email (si fourni) - vérification exacte
    2. Nom + Prénom + Téléphone (si tous fournis) - vérification exacte
    3. Nom + Prénom (si fournis) - vérification exacte (moins fiable)
    
    Retourne le candidat existant si trouvé, None sinon
    """
    # Vérification par email (le plus fiable)
    if email and email.strip():
        email_clean = email.strip().lower()
        existing = session.exec(
            select(Candidate).where(
                Candidate.email.isnot(None),
                func.lower(Candidate.email) == email_clean
            )
        ).first()
        if existing:
            return existing
    
    # Vérification par nom + prénom + téléphone (si tous fournis)
    if first_name and last_name and phone:
        first_name_clean = first_name.strip().lower()
        last_name_clean = last_name.strip().lower()
        phone_clean = phone.strip()
        
        existing = session.exec(
            select(Candidate).where(
                func.lower(Candidate.first_name) == first_name_clean,
                func.lower(Candidate.last_name) == last_name_clean,
                Candidate.phone.isnot(None),
                Candidate.phone == phone_clean
            )
        ).first()
        if existing:
            return existing
    
    # Vérification par nom + prénom uniquement (moins fiable, mais utile si pas d'email/téléphone)
    if first_name and last_name:
        first_name_clean = first_name.strip().lower()
        last_name_clean = last_name.strip().lower()
        
        existing = session.exec(
            select(Candidate).where(
                func.lower(Candidate.first_name) == first_name_clean,
                func.lower(Candidate.last_name) == last_name_clean
            )
        ).first()
        if existing:
            return existing
    
    return None


def extract_text_from_pdf(file_path: str) -> str:
    """Extrait le texte d'un fichier PDF"""
    if fitz is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PyMuPDF n'est pas installé. Installez-le avec: pip install pymupdf"
        )
    
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Erreur lors de l'extraction du PDF: {str(e)}"
        )


def extract_text_from_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier Word (.docx)"""
    if Document is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="python-docx n'est pas installé. Installez-le avec: pip install python-docx"
        )
    
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Erreur lors de l'extraction du document Word: {str(e)}"
        )


def extract_image_from_pdf(file_path: str) -> Optional[str]:
    """Extrait la première image d'un fichier PDF (généralement la photo de profil)"""
    if fitz is None:
        return None
    
    try:
        doc = fitz.open(file_path)
        # Chercher la première image sur la première page (généralement où se trouve la photo)
        for page_num in range(min(2, len(doc))):  # Vérifier les 2 premières pages
            page = doc[page_num]
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Convertir en base64
                    image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    # Retourner avec le préfixe data URI
                    mime_type = f"image/{image_ext}" if image_ext != "jpg" else "image/jpeg"
                    doc.close()
                    return f"data:{mime_type};base64,{image_base64}"
                except Exception as e:
                    continue
        
        doc.close()
        return None
    except Exception as e:
        return None


def extract_image_from_docx(file_path: str) -> Optional[str]:
    """Extrait la première image d'un document Word"""
    if Document is None:
        return None
    
    try:
        doc = Document(file_path)
        
        # Parcourir les relations du document pour trouver les images
        # Les images dans Word sont stockées dans les relations
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # Convertir en base64
                    image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    # Déterminer le type MIME
                    content_type = image_part.content_type
                    return f"data:{content_type};base64,{image_base64}"
                except Exception as e:
                    continue
        
        return None
    except Exception as e:
        return None


def extract_image_from_cv(file_path: str, file_extension: str) -> Optional[str]:
    """Extrait la première image d'un CV (PDF ou Word)"""
    try:
        if file_extension == ".pdf":
            return extract_image_from_pdf(file_path)
        elif file_extension in {".doc", ".docx"}:
            return extract_image_from_docx(file_path)
        else:
            return None
    except Exception as e:
        return None


async def extract_text_from_cv(file: UploadFile) -> tuple[str, str]:
    """Extrait le texte brut d'un CV (PDF ou Word) et retourne aussi le chemin temporaire"""
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Aucun nom de fichier fourni"
        )
    
    file_extension = Path(file.filename).suffix.lower()
    
    # Créer un fichier temporaire
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        tmp_path = tmp_file.name
        try:
            # Réinitialiser le pointeur du fichier
            await file.seek(0)
            
            # Lire le contenu du fichier uploadé
            content = await file.read()
            
            # Écrire le contenu dans le fichier temporaire
            tmp_file.write(content)
            tmp_file.flush()
            
            # Extraire le texte selon le type de fichier
            if file_extension == ".pdf":
                text = extract_text_from_pdf(tmp_path)
            elif file_extension in {".doc", ".docx"}:
                text = extract_text_from_docx(tmp_path)
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Format de fichier non supporté: {file_extension}"
                )
            
            return text, tmp_path
        except HTTPException:
            # Nettoyer le fichier temporaire en cas d'erreur
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
            raise
        except Exception as e:
            # Nettoyer le fichier temporaire en cas d'erreur
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Erreur lors de l'extraction du texte du CV: {str(e)}"
            )


def parse_cv_with_llm(cv_text: str) -> dict:
    """Utilise un LLM pour parser le texte du CV et extraire les informations structurées"""
    if genai is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="google-generativeai n'est pas installé. Installez-le avec: pip install google-generativeai"
        )
    
    # Récupérer la clé API depuis les variables d'environnement
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("OPENAI_API_KEY")  # Compatibilité avec ancienne variable
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="GEMINI_API_KEY n'est pas configurée dans les variables d'environnement"
        )
    
    # Configurer Gemini
    genai.configure(api_key=api_key)
    # Utiliser gemini-1.5-pro (modèle stable et disponible)
    model = genai.GenerativeModel('gemini-1.5-pro')
    
    # Prompt pour extraire les informations du CV
    prompt = f"""Tu es un assistant expert en recrutement. Analyse le CV suivant et extrais les informations pertinentes au format JSON.

CV:
{cv_text}

Extrais les informations suivantes au format JSON strict (sans commentaires, sans markdown):
{{
  "first_name": "Prénom du candidat",
  "last_name": "Nom du candidat",
  "profile_title": "Titre du poste actuel ou recherché (ex: Développeur Fullstack, Chef de projet, etc.)",
  "years_of_experience": nombre d'années d'expérience total (entier, 0 si débutant),
  "email": "Email du candidat",
  "phone": "Téléphone du candidat",
  "skills": ["compétence1", "compétence2", ...],
  "source": "Source du CV si mentionnée (LinkedIn, APEC, etc.)",
  "notes": "Notes pertinentes extraites du CV"
}}

Règles importantes:
- Si une information n'est pas trouvée, utilise null pour les champs optionnels
- first_name et last_name sont obligatoires (extrais-les du nom complet)
- years_of_experience doit être un nombre entier (calcule-le à partir des dates d'expérience)
- skills doit être une liste de chaînes (extrais les technologies, langages, outils mentionnés)
- Retourne UNIQUEMENT le JSON, sans texte avant ou après
"""
    
    try:
        # Construire le prompt complet pour Gemini
        full_prompt = f"""Tu es un assistant expert en extraction de données de CV. Tu retournes uniquement du JSON valide.

{prompt}"""
        
        # Configuration pour la génération
        generation_config = {
            "temperature": 0.1,
            "max_output_tokens": 2000,
        }
        
        response = model.generate_content(
            full_prompt,
            generation_config=generation_config
        )
        
        # Extraire le JSON de la réponse
        response_text = response.text.strip()
        
        # Nettoyer la réponse (enlever les markdown code blocks si présents)
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        # Parser le JSON
        parsed_data = json.loads(response_text)
        
        # Valider que first_name et last_name sont présents
        if not parsed_data.get("first_name") or not parsed_data.get("last_name"):
            raise ValueError("first_name et last_name sont obligatoires")
        
        return parsed_data
        
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors du parsing JSON de la réponse LLM: {str(e)}. Réponse reçue: {response_text[:200]}"
        )
    except Exception as e:
        # Vérifier si c'est une erreur de rate limit (429)
        error_str = str(e)
        is_rate_limit = False
        
        # Vérifier si c'est une erreur de rate limit Gemini
        if GeminiError and GeminiError != Exception and isinstance(e, GeminiError):
            is_rate_limit = "429" in error_str or "quota" in error_str.lower() or "rate_limit" in error_str.lower()
        elif "429" in error_str or "rate_limit" in error_str.lower() or "Rate limit" in error_str or "quota" in error_str.lower() or "rate_limit_exceeded" in error_str.lower():
            is_rate_limit = True
        
        if is_rate_limit:
            # Extraire le temps d'attente si disponible
            wait_time = None
            if "try again in" in error_str:
                try:
                    import re
                    # Chercher différents formats de temps (19h16m53.76s, 1h, 30m, etc.)
                    match = re.search(r'try again in ([\d\.]+[hms]+)', error_str)
                    if match:
                        wait_time = match.group(1)
                except:
                    pass
            
            error_message = "Limite de requêtes Gemini API atteinte. Le service d'analyse automatique de CV est temporairement indisponible."
            if wait_time:
                error_message += f" Veuillez réessayer dans {wait_time}."
            else:
                error_message += " Veuillez réessayer plus tard."
            
            error_message += " Vous pouvez créer le candidat manuellement en remplissant le formulaire."
            
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=error_message
            )
        
        # Autres erreurs Gemini ou erreurs générales
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de l'appel à l'API Gemini: {str(e)}"
        )
    except Exception as e:
        # Gérer les autres exceptions (y compris les erreurs Gemini non capturées)
        error_str = str(e)
        if "429" in error_str or "rate_limit" in error_str.lower() or "Rate limit" in error_str:
            # Extraire le temps d'attente si disponible
            wait_time = None
            if "try again in" in error_str:
                try:
                    import re
                    match = re.search(r'try again in ([\d\.]+[hms]+)', error_str)
                    if match:
                        wait_time = match.group(1)
                except:
                    pass
            
            error_message = "Limite de requêtes Gemini API atteinte. Le service d'analyse automatique de CV est temporairement indisponible."
            if wait_time:
                error_message += f" Veuillez réessayer dans {wait_time}."
            else:
                error_message += " Veuillez réessayer plus tard."
            
            error_message += " Vous pouvez créer le candidat manuellement en remplissant le formulaire."
            
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=error_message
            )
        
        # Autres erreurs
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de l'appel à l'API Gemini: {str(e)}"
        )


@router.post("/upload-photo", status_code=status.HTTP_200_OK)
async def upload_candidate_photo(
    photo: UploadFile = File(...),
    current_user: User = Depends(require_recruteur),
):
    """
    Upload une photo de profil pour un candidat
    
    Retourne l'URL de la photo uploadée.
    """
    if not photo.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Aucun fichier fourni"
        )
    
    if not is_allowed_image(photo.filename):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Format d'image non autorisé. Formats acceptés: {', '.join(ALLOWED_IMAGE_EXTENSIONS)}"
        )
    
    # Vérifier que le dossier de destination existe (sécurité bonus)
    PHOTOS_DIR.mkdir(parents=True, exist_ok=True)
    
    # Générer un nom de fichier unique
    file_extension = Path(photo.filename).suffix
    unique_filename = f"{uuid4().hex}{file_extension}"
    file_path = PHOTOS_DIR / unique_filename
    
    # Sauvegarder le fichier
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(photo.file, buffer)
    
    # Retourner l'URL relative (sera servie par FastAPI via /static)
    photo_url = f"/static/uploads/{unique_filename}"
    
    return {"photo_url": photo_url, "filename": unique_filename}


@router.post("/parse-cv", response_model=CandidateParseResponse)
async def parse_cv(
    cv_file: UploadFile = File(..., description="Fichier CV (PDF ou Word)"),
    current_user: User = Depends(require_recruteur),
):
    """
    Parse un CV et extrait automatiquement les informations du candidat, y compris la photo
    
    Accepte un fichier PDF ou Word, extrait le texte et les images, et utilise un LLM
    pour structurer les données selon le modèle CandidateCreate.
    """
    tmp_path = None
    try:
        # Vérifier que le fichier est autorisé
        if not is_allowed_file(cv_file.filename):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Format de fichier non supporté. Formats acceptés: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        file_extension = Path(cv_file.filename).suffix.lower()
        
        # Vérifier que le fichier a un nom
        if not cv_file.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Aucun nom de fichier fourni"
            )
        
        # Réinitialiser le pointeur du fichier pour pouvoir le lire plusieurs fois
        await cv_file.seek(0)
        
        # Extraire le texte du CV (retourne aussi le chemin temporaire)
        cv_text, tmp_path = await extract_text_from_cv(cv_file)
        
        if not cv_text or len(cv_text.strip()) < 50:
            # Nettoyer le fichier temporaire
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Le CV semble vide ou le texte n'a pas pu être extrait correctement"
            )
        
        # Extraire l'image du CV
        profile_picture_base64 = None
        if tmp_path and os.path.exists(tmp_path):
            try:
                profile_picture_base64 = extract_image_from_cv(tmp_path, file_extension)
            except Exception as e:
                # Si l'extraction d'image échoue, continuer sans image
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f"Erreur lors de l'extraction de l'image du CV: {str(e)}")
        
        # Parser le texte avec le LLM
        parsed_data = parse_cv_with_llm(cv_text)
        
        # Valider et retourner les données
        # Convertir years_of_experience en int si présent
        if "years_of_experience" in parsed_data and parsed_data["years_of_experience"] is not None:
            try:
                parsed_data["years_of_experience"] = int(parsed_data["years_of_experience"])
            except (ValueError, TypeError):
                parsed_data["years_of_experience"] = None
        
        # S'assurer que skills est une liste
        if "skills" in parsed_data and parsed_data["skills"] is None:
            parsed_data["skills"] = []
        elif "skills" not in parsed_data:
            parsed_data["skills"] = []
        
        # S'assurer que tags est une liste (peut être vide)
        if "tags" not in parsed_data:
            parsed_data["tags"] = []
        
        # Créer la réponse avec le schéma CandidateParseResponse
        response_data = CandidateParseResponse(
            first_name=parsed_data.get("first_name", ""),
            last_name=parsed_data.get("last_name", ""),
            profile_title=parsed_data.get("profile_title"),
            years_of_experience=parsed_data.get("years_of_experience"),
            email=parsed_data.get("email"),
            phone=parsed_data.get("phone"),
            tags=parsed_data.get("tags", []),
            skills=parsed_data.get("skills", []),
            source=parsed_data.get("source"),
            notes=parsed_data.get("notes"),
            profile_picture_base64=profile_picture_base64
        )
        
        return response_data
        
    except HTTPException:
        # Nettoyer le fichier temporaire en cas d'erreur HTTP
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise
    except Exception as e:
        # Nettoyer le fichier temporaire en cas d'erreur
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Erreur lors du parsing du CV: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors du parsing du CV: {str(e)}"
        )
    finally:
        # S'assurer que le fichier temporaire est toujours nettoyé
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


@router.post("/", response_model=CandidateResponse, status_code=status.HTTP_201_CREATED)
async def create_candidate(
    first_name: str = Form(...),
    last_name: str = Form(...),
    profile_title: Optional[str] = Form(None),  # Titre du profil
    years_of_experience: Optional[int] = Form(None),  # Années d'expérience
    email: Optional[str] = Form(None),
    phone: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),  # Reçu comme string séparée par des virgules
    skills: Optional[str] = Form(None),  # Reçu comme string séparée par des virgules
    profile_picture_url: Optional[str] = Form(None),  # URL de la photo uploadée
    source: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    cv_file: Optional[UploadFile] = File(None),
    current_user: User = Depends(require_recruteur),
    session: Session = Depends(get_session)
):
    """
    Créer un nouveau candidat (US04)
    
    Permet de créer une fiche candidat avec CV et tags.
    """
    # Utiliser l'utilisateur connecté
    created_by = current_user.id
    
    # Traiter les tags (string séparée par virgules -> liste)
    tags_list = None
    if tags:
        tags_list = [tag.strip() for tag in tags.split(",") if tag.strip()]
        # Si la liste est vide après nettoyage, passer None
        if not tags_list:
            tags_list = None
    
    # Traiter les compétences (string séparée par virgules -> liste pour PostgreSQL ARRAY)
    # Les compétences sont stockées comme PostgreSQL ARRAY (TEXT[])
    skills_list = None
    if skills:
        # Nettoyer et valider les compétences
        skills_cleaned = [skill.strip() for skill in skills.split(",") if skill.strip()]
        if skills_cleaned:
            skills_list = skills_cleaned  # Stocker comme liste pour PostgreSQL ARRAY
    
    # Gérer l'upload du CV
    cv_file_path = None
    if cv_file and cv_file.filename:
        if not is_allowed_file(cv_file.filename):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Format de fichier non autorisé. Formats acceptés: PDF, DOC, DOCX"
            )
        
        # Générer un nom de fichier unique
        file_extension = Path(cv_file.filename).suffix
        unique_filename = f"{uuid4().hex}{file_extension}"
        cv_file_path = str(UPLOAD_DIR / unique_filename)
        
        # Sauvegarder le fichier
        with open(cv_file_path, "wb") as buffer:
            shutil.copyfileobj(cv_file.file, buffer)
    
    # Vérifier les doublons avant de créer le candidat
    existing_candidate = check_duplicate_candidate(
        session=session,
        email=email.strip() if email else None,
        first_name=first_name.strip(),
        last_name=last_name.strip(),
        phone=phone.strip() if phone else None
    )
    
    if existing_candidate:
        # Si un candidat existe déjà, retourner une erreur avec les informations du candidat existant
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={
                "message": "Un candidat avec ces informations existe déjà dans la plateforme",
                "existing_candidate_id": str(existing_candidate.id),
                "existing_candidate_name": f"{existing_candidate.first_name} {existing_candidate.last_name}",
                "existing_candidate_email": existing_candidate.email,
                "match_criteria": "email" if email and email.strip() else ("nom + prénom + téléphone" if phone and phone.strip() else "nom + prénom")
            }
        )
    
    # Créer le candidat
    try:
        candidate = Candidate(
            first_name=first_name.strip(),
            last_name=last_name.strip(),
            profile_title=profile_title.strip() if profile_title else None,
            years_of_experience=years_of_experience,
            email=email.strip() if email else None,
            phone=phone.strip() if phone else None,
            cv_file_path=cv_file_path,
            profile_picture_url=profile_picture_url.strip() if profile_picture_url else None,
            # photo_url n'est pas stocké en DB, on utilise seulement profile_picture_url
            tags=tags_list if tags_list else None,  # S'assurer que None est utilisé si la liste est vide
            skills=skills_list if skills_list else [],  # Compétences stockées comme liste PostgreSQL ARRAY (défaut: [])
            source=source.strip() if source else None,
            status="sourcé",  # Statut par défaut selon la contrainte CHECK de la base de données
            notes=notes.strip() if notes else None,
            created_by=created_by
        )
        
        session.add(candidate)
        session.commit()
        session.refresh(candidate)
        
        # Normaliser la réponse comme dans get_candidate et list_candidates
        # Créer un dictionnaire avec toutes les valeurs
        creator_info = get_creator_info(candidate.created_by, session)
        candidate_dict = {
            "id": candidate.id,
            "first_name": candidate.first_name,
            "last_name": candidate.last_name,
            "profile_title": candidate.profile_title,
            "years_of_experience": candidate.years_of_experience,
            "email": candidate.email,
            "phone": candidate.phone,
            "cv_file_path": candidate.cv_file_path,
            "profile_picture_url": candidate.profile_picture_url,
            "photo_url": candidate.profile_picture_url,  # Alias de profile_picture_url (non mappé en DB)
            "tags": candidate.tags if candidate.tags else None,
            "skills": candidate.skills if candidate.skills else [],  # Convertir None en []
            "source": candidate.source,
            "status": candidate.status,
            "notes": candidate.notes,
            "created_by": candidate.created_by,
            "created_at": candidate.created_at,
            "updated_at": candidate.updated_at,
            **creator_info,
        }
        
        # Valider avec le schéma Pydantic
        return CandidateResponse.model_validate(candidate_dict)
    except Exception as e:
        session.rollback()
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Erreur lors de la création du candidat: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de la création du candidat: {str(e)}"
        )


@router.get("/", response_model=List[CandidateResponse])
def list_candidates(
    skip: int = 0,
    limit: int = 100,
    tag_filter: Optional[str] = Query(None, description="Filtrer par tag"),
    source_filter: Optional[str] = Query(None, description="Filtrer par source"),
    status_filter: Optional[str] = Query(None, description="Filtrer par statut"),
    current_user: User = Depends(get_current_active_user),
    session: Session = Depends(get_session)
):
    """
    Lister tous les candidats avec filtres optionnels
    
    Règles d'accès:
    - Client: Ne voit que les candidats en shortlist pour ses propres postes
    - Recruteur/Manager/Admin: Voit tous les candidats
    
    ⚠️ IMPORTANT : 
    - Le champ 'photo_url' n'existe PAS dans la base de données, c'est juste un alias de 'profile_picture_url' dans le schéma de réponse.
    - Si vous obtenez une erreur concernant 'profile_picture_url' ou 'skills', exécutez la migration SQL :
      psql -U postgres -d recrutement_db -c "ALTER TABLE candidates ADD COLUMN IF NOT EXISTS profile_picture_url VARCHAR(500); ALTER TABLE candidates ADD COLUMN IF NOT EXISTS skills TEXT[];"
    """
    try:
        # IMPORTANT: select(Candidate) ne charge que les colonnes définies dans le modèle Candidate
        # Le modèle n'a PAS de champ 'photo_url', seulement 'profile_picture_url'
        # 'photo_url' est un alias ajouté dans le schéma de réponse (CandidateResponse)
        statement = select(Candidate)
        
        # Règle d'accès: Les clients ne voient que les candidats en shortlist pour leurs postes
        user_role = current_user.role if isinstance(current_user.role, str) else current_user.role.value
        if user_role == UserRole.CLIENT.value:
            # Trouver les jobs du client
            from models import Job
            from sqlalchemy import or_
            client_jobs_statement = select(Job).where(Job.department == current_user.department)
            client_jobs = session.exec(client_jobs_statement).all()
            job_ids = [job.id for job in client_jobs]
            
            if not job_ids:
                return []
            
            # Trouver les applications en shortlist pour ces jobs
            applications_statement = select(Application).where(
                Application.job_id.in_(job_ids)  # type: ignore
            ).where(
                Application.is_in_shortlist == True
            )
            applications = session.exec(applications_statement).all()
            candidate_ids = [app.candidate_id for app in applications]
            
            if not candidate_ids:
                return []
            
            # Filtrer par IDs de candidats
            statement = statement.where(
                or_(*[Candidate.id == cid for cid in candidate_ids])
            )
            # Forcer le statut à shortlist pour les clients
            statement = statement.where(Candidate.status == "shortlist")
        else:
            # Recruteurs, Managers, Admins: peuvent voir tous les candidats
            # Filtre par source
            if source_filter:
                statement = statement.where(Candidate.source == source_filter)
            
            # Filtre par statut
            if status_filter:
                statement = statement.where(Candidate.status == status_filter)
        
        statement = statement.offset(skip).limit(limit).order_by(Candidate.created_at.desc())
        
        candidates = session.exec(statement).all()
        
        # Convertir explicitement en CandidateResponse pour éviter les problèmes de sérialisation
        from schemas import CandidateResponse
        candidates_list = []
        for candidate in candidates:
            try:
                # Récupérer les informations du créateur
                creator_info = get_creator_info(candidate.created_by, session)
                # Créer un dictionnaire avec tous les champs nécessaires
                candidate_dict = {
                    "id": candidate.id,
                    "first_name": candidate.first_name,
                    "last_name": candidate.last_name,
                    "profile_title": candidate.profile_title,
                    "years_of_experience": candidate.years_of_experience,
                    "email": candidate.email,
                    "phone": candidate.phone,
                    "cv_file_path": candidate.cv_file_path,
                    "profile_picture_url": candidate.profile_picture_url,
                    "photo_url": candidate.profile_picture_url,  # Alias pour compatibilité
                    "tags": candidate.tags or [],
                    "skills": candidate.skills or [],
                    "source": candidate.source,
                    "status": candidate.status,
                    "notes": candidate.notes,
                    "created_by": candidate.created_by,
                    "created_at": candidate.created_at,
                    "updated_at": candidate.updated_at,
                    **creator_info,
                }
                candidates_list.append(CandidateResponse.model_validate(candidate_dict))
            except Exception as e:
                logger.warning(f"Erreur lors de la conversion du candidat {candidate.id}: {e}")
                continue
        
        # Logs de debug pour voir ce que la base de données renvoie
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"🔍 [DEBUG] Nombre de candidats récupérés: {len(candidates_list)}")
        if candidates_list:
            first_candidate = candidates_list[0]
            logger.info(f"🔍 [DEBUG] Premier candidat - ID: {first_candidate.id}")
            logger.info(f"🔍 [DEBUG] Premier candidat - Nom: {first_candidate.first_name} {first_candidate.last_name}")
            logger.info(f"🔍 [DEBUG] Premier candidat - profile_picture_url: {first_candidate.profile_picture_url}")
            logger.info(f"🔍 [DEBUG] Premier candidat - skills: {first_candidate.skills}")
            logger.info(f"🔍 [DEBUG] Premier candidat - tags: {first_candidate.tags}")
            logger.info(f"🔍 [DEBUG] Premier candidat - status: {first_candidate.status}")
        
        # Filtrer par tag en Python si nécessaire (pour éviter les problèmes SQL)
        if tag_filter and user_role != UserRole.CLIENT.value:
            candidates_list = [
                c for c in candidates_list 
                if c.tags and tag_filter in c.tags
            ]
        
        return candidates_list
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        error_msg = str(e)
        error_type = type(e).__name__
        
        # ROLLBACK IMPÉRATIF : La session est dans un état invalide après l'erreur SQL
        # DOIT être fait en PREMIER, avant toute autre opération
        session.rollback()
        logger.info("✅ [TRANSACTION] Session rollback effectué après erreur")
        
        logger.error(f"❌ [ERREUR 500] Erreur lors de la récupération des candidats: {error_type}: {error_msg}", exc_info=True)
        
        # Détecter si c'est une erreur de colonne manquante
        if "photo_url" in error_msg or "profile_picture_url" in error_msg or "skills" in error_msg or "does not exist" in error_msg or "UndefinedColumn" in error_msg:
            # Utiliser une requête SQL brute qui ne sélectionne que les colonnes existantes
            # IMPORTANT: photo_url n'existe PAS dans la base de données, c'est juste un alias dans le schéma de réponse
            # On essaie d'abord avec profile_picture_url, sinon on l'exclut
            try:
                from sqlalchemy import text as sql_text
                
                # Vérifier si profile_picture_url existe dans la table
                # On essaie d'abord une requête simple pour détecter les colonnes disponibles
                check_columns_query = """
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name = 'candidates' 
                    AND column_name IN ('profile_picture_url', 'skills')
                """
                try:
                    check_result = session.exec(sql_text(check_columns_query))
                    available_columns = {row[0] for row in check_result.all()}
                except:
                    available_columns = set()
                
                # Construire la requête SQL en sélectionnant uniquement les colonnes qui existent
                # Colonnes de base toujours présentes:
                base_columns = "id, first_name, last_name, email, phone, cv_file_path, tags, source, status, notes, created_by, created_at, updated_at"
                
                # Ajouter profile_picture_url seulement s'il existe
                if 'profile_picture_url' in available_columns:
                    sql_query = f"""
                        SELECT {base_columns}, profile_picture_url
                        FROM candidates
                        WHERE 1=1
                    """
                    has_profile_picture = True
                else:
                    sql_query = f"""
                        SELECT {base_columns}
                        FROM candidates
                        WHERE 1=1
                    """
                    has_profile_picture = False
                
                # Ajouter skills seulement s'il existe (mais on ne l'utilise pas dans le fallback pour simplifier)
                has_skills = 'skills' in available_columns
                
                user_role = current_user.role if isinstance(current_user.role, str) else current_user.role.value
                
                if user_role == UserRole.CLIENT.value:
                    from models import Job, Application
                    from sqlalchemy import or_
                    client_jobs_statement = select(Job).where(Job.department == current_user.department)
                    client_jobs = session.exec(client_jobs_statement).all()
                    job_ids = [job.id for job in client_jobs]
                    
                    if not job_ids:
                        return []
                    
                    applications_statement = select(Application).where(
                        Application.job_id.in_(job_ids)  # type: ignore
                    ).where(Application.is_in_shortlist == True)
                    applications = session.exec(applications_statement).all()
                    candidate_ids = [app.candidate_id for app in applications]
                    
                    if not candidate_ids:
                        return []
                    
                    uuid_list = ','.join([f"'{str(cid)}'" for cid in candidate_ids])
                    sql_query += f" AND id = ANY(ARRAY[{uuid_list}]::UUID[]) AND status = 'shortlist'"
                else:
                    if source_filter:
                        sql_query += f" AND source = '{source_filter.replace(chr(39), chr(39)+chr(39))}'"
                    if status_filter:
                        sql_query += f" AND status = '{status_filter.replace(chr(39), chr(39)+chr(39))}'"
                
                sql_query += f" ORDER BY created_at DESC LIMIT {limit} OFFSET {skip}"
                
                result = session.exec(sql_text(sql_query))
                rows = result.all()
                
                # Convertir en dictionnaires (CandidateResponse)
                # L'ordre des colonnes dépend de si profile_picture_url existe ou non
                # Sans profile_picture_url: 0: id, 1: first_name, 2: last_name, 3: email, 4: phone, 5: cv_file_path,
                # 6: tags, 7: source, 8: status, 9: notes, 10: created_by, 11: created_at, 12: updated_at
                # Avec profile_picture_url: même ordre mais profile_picture_url à l'index 13
                candidates_response = []
                for row in rows:
                    try:
                        row_index = 0
                        candidate_dict = {
                            'id': str(row[row_index]) if len(row) > row_index and row[row_index] else None,
                            'first_name': row[row_index + 1] if len(row) > row_index + 1 else '',
                            'last_name': row[row_index + 2] if len(row) > row_index + 2 else '',
                            'profile_title': None,  # Colonne manquante dans le fallback SQL
                            'years_of_experience': None,  # Colonne manquante dans le fallback SQL
                            'email': row[row_index + 3] if len(row) > row_index + 3 else None,
                            'phone': row[row_index + 4] if len(row) > row_index + 4 else None,
                            'cv_file_path': row[row_index + 5] if len(row) > row_index + 5 else None,
                            'tags': row[row_index + 6] if len(row) > row_index + 6 and row[row_index + 6] else None,
                            'source': row[row_index + 7] if len(row) > row_index + 7 else None,
                            'status': row[row_index + 8] if len(row) > row_index + 8 else 'sourcé',
                            'notes': row[row_index + 9] if len(row) > row_index + 9 else None,
                            'created_by': str(row[row_index + 10]) if len(row) > row_index + 10 and row[row_index + 10] else None,
                            'created_at': row[row_index + 11] if len(row) > row_index + 11 else None,
                            'updated_at': row[row_index + 12] if len(row) > row_index + 12 else None,
                        }
                        
                        # Ajouter profile_picture_url si disponible
                        if has_profile_picture and len(row) > row_index + 13:
                            profile_pic = row[row_index + 13]
                            candidate_dict['profile_picture_url'] = profile_pic
                            candidate_dict['photo_url'] = profile_pic  # Alias
                        else:
                            candidate_dict['profile_picture_url'] = None
                            candidate_dict['photo_url'] = None  # photo_url n'existe jamais en DB, c'est juste un alias
                        
                        # Ajouter skills (toujours None dans le fallback pour simplifier)
                        candidate_dict['skills'] = []
                        # Log de debug pour la première ligne
                        if len(candidates_response) == 0:
                            logger.info(f"🔍 [DEBUG SQL Fallback] Premier candidat converti: {candidate_dict}")
                        candidates_response.append(candidate_dict)
                    except Exception as row_error:
                        logger.warning(f"Erreur lors de la conversion d'une ligne: {row_error}, row length: {len(row) if hasattr(row, '__len__') else 'N/A'}")
                        continue
                
                # Filtrer par tag si nécessaire
                if tag_filter and user_role != UserRole.CLIENT.value:
                    candidates_response = [
                        c for c in candidates_response 
                        if c.get('tags') and tag_filter in c['tags']
                    ]
                
                return candidates_response
                
            except Exception as sql_error:
                logger.error(f"Erreur avec la requête SQL brute: {str(sql_error)}", exc_info=True)
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=(
                        "Erreur de base de données : Les colonnes 'profile_picture_url' et 'skills' sont manquantes. "
                        "Exécutez la migration SQL : "
                        "psql -U postgres -d recrutement_db -c \"ALTER TABLE candidates ADD COLUMN IF NOT EXISTS profile_picture_url VARCHAR(500); ALTER TABLE candidates ADD COLUMN IF NOT EXISTS skills TEXT[];\""
                    )
                )
        
        # Si ce n'est pas une erreur de colonne manquante, relancer l'erreur originale
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de la récupération des candidats: {error_msg}"
        )


# Routes spécifiques AVANT les routes génériques pour éviter les conflits de routage
@router.patch("/{candidate_id}/status", response_model=CandidateResponse)
def update_candidate_status(
    candidate_id: UUID,
    new_status: str = Query(..., description="Nouveau statut"),
    current_user: User = Depends(require_recruteur),
    session: Session = Depends(get_session)
):
    """
    Mettre à jour le statut d'un candidat
    
    Règle métier : Pour passer un candidat à "shortlist", il faut qu'il y ait au moins un entretien avec un feedback.
    """
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f"🔍 [DEBUG] update_candidate_status appelé - candidate_id: {candidate_id}, new_status: {new_status}")
    
    candidate = session.get(Candidate, candidate_id)
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidat non trouvé"
        )
    
    # Vérifier si le changement de statut nécessite un feedback
    if new_status in ["shortlist", "offre"] and candidate.status not in ["shortlist", "offre"]:
        # Trouver les applications de ce candidat
        applications_statement = select(Application).where(Application.candidate_id == candidate_id)
        applications = session.exec(applications_statement).all()
        
        if not applications:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Veuillez saisir un feedback avant de changer le statut"
            )
        
        # Vérifier qu'il existe au moins un entretien avec un feedback non vide pour ce candidat
        has_feedback = False
        for app in applications:
            # Vérifier les entretiens avec feedback non null et non vide
            interviews_statement = select(Interview).where(
                Interview.application_id == app.id
            )
            all_interviews = session.exec(interviews_statement).all()
            
            for interview in all_interviews:
                if interview.feedback and interview.feedback.strip():
                    has_feedback = True
                    break
            
            if has_feedback:
                break
        
        if not has_feedback:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Veuillez saisir un feedback avant de changer le statut"
            )
    
    candidate.status = new_status
    session.add(candidate)
    session.commit()
    session.refresh(candidate)
    
    # Normaliser la réponse comme dans les autres endpoints
    creator_info = get_creator_info(candidate.created_by, session)
    candidate_dict = {
        "id": candidate.id,
        "first_name": candidate.first_name,
        "last_name": candidate.last_name,
        "profile_title": candidate.profile_title,
        "years_of_experience": candidate.years_of_experience,
        "email": candidate.email,
        "phone": candidate.phone,
        "cv_file_path": candidate.cv_file_path,
        "profile_picture_url": candidate.profile_picture_url,
        "photo_url": candidate.profile_picture_url,  # Alias de profile_picture_url (non mappé en DB)
        "tags": candidate.tags if candidate.tags else None,
        "skills": candidate.skills if candidate.skills else [],  # Convertir None en []
        "source": candidate.source,
        "status": candidate.status,
        "notes": candidate.notes,
        "created_by": candidate.created_by,
        "created_at": candidate.created_at,
        "updated_at": candidate.updated_at,
        **creator_info,
    }
    
    # Valider avec le schéma Pydantic
    return CandidateResponse.model_validate(candidate_dict)


@router.get("/{candidate_id}", response_model=CandidateResponse)
def get_candidate(
    candidate_id: UUID,
    current_user: User = Depends(get_current_active_user),
    session: Session = Depends(get_session)
):
    """
    Récupérer un candidat par son ID
    
    Accessible à tous les utilisateurs authentifiés (Recruteur, Manager, Client, Administrateur).
    Normalise la réponse pour inclure photo_url (alias de profile_picture_url) et tous les champs,
    y compris profile_title et years_of_experience.
    """
    try:
        candidate = session.get(Candidate, candidate_id)
        if not candidate:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Candidat non trouvé"
            )
        
        # Normaliser la réponse comme dans list_candidates
        # Créer un dictionnaire avec toutes les valeurs
        creator_info = get_creator_info(candidate.created_by, session)
        candidate_dict = {
            "id": candidate.id,
            "first_name": candidate.first_name,
            "last_name": candidate.last_name,
            "profile_title": candidate.profile_title,
            "years_of_experience": candidate.years_of_experience,
            "email": candidate.email,
            "phone": candidate.phone,
            "cv_file_path": candidate.cv_file_path,
            "profile_picture_url": candidate.profile_picture_url,
            "photo_url": candidate.profile_picture_url,  # Alias de profile_picture_url (non mappé en DB)
            "tags": candidate.tags if candidate.tags else None,
            "skills": candidate.skills if candidate.skills else [],  # Convertir None en []
            "source": candidate.source,
            "status": candidate.status,
            "notes": candidate.notes,
            "created_by": candidate.created_by,
            "created_at": candidate.created_at,
            "updated_at": candidate.updated_at,
            **creator_info,
        }
        
        # Valider avec le schéma Pydantic
        return CandidateResponse.model_validate(candidate_dict)
        
    except HTTPException:
        raise
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Erreur lors de la récupération du candidat {candidate_id}: {str(e)}", exc_info=True)
        
        # Rollback en cas d'erreur
        session.rollback()
        
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de la récupération du candidat: {str(e)}"
        )


@router.patch("/{candidate_id}", response_model=CandidateResponse)
def update_candidate(
    candidate_id: UUID,
    candidate_update: CandidateUpdate,
    current_user: User = Depends(get_current_active_user),
    session: Session = Depends(get_session)
):
    """
    Mettre à jour un candidat
    """
    candidate = session.get(Candidate, candidate_id)
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidat non trouvé"
        )
    
    # Mettre à jour les champs fournis
    update_data = candidate_update.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(candidate, field, value)
    
    session.add(candidate)
    session.commit()
    session.refresh(candidate)
    
    # Normaliser la réponse comme dans get_candidate et create_candidate
    creator_info = get_creator_info(candidate.created_by, session)
    candidate_dict = {
        "id": candidate.id,
        "first_name": candidate.first_name,
        "last_name": candidate.last_name,
        "profile_title": candidate.profile_title,
        "years_of_experience": candidate.years_of_experience,
        "email": candidate.email,
        "phone": candidate.phone,
        "cv_file_path": candidate.cv_file_path,
        "profile_picture_url": candidate.profile_picture_url,
        "photo_url": candidate.profile_picture_url,  # Alias de profile_picture_url (non mappé en DB)
        "tags": candidate.tags if candidate.tags else None,
        "skills": candidate.skills if candidate.skills else [],  # Convertir None en []
        "source": candidate.source,
        "status": candidate.status,
        "notes": candidate.notes,
        "created_by": candidate.created_by,
        "created_at": candidate.created_at,
        "updated_at": candidate.updated_at,
        **creator_info,
    }
    
    # Valider avec le schéma Pydantic
    return CandidateResponse.model_validate(candidate_dict)


@router.get("/{candidate_id}/cv", response_class=FileResponse)
def download_cv(
    candidate_id: UUID,
    session: Session = Depends(get_session)
):
    """
    Télécharger le CV d'un candidat
    """
    candidate = session.get(Candidate, candidate_id)
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidat non trouvé"
        )
    
    if not candidate.cv_file_path or not os.path.exists(candidate.cv_file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV non trouvé"
        )
    
    return FileResponse(
        path=candidate.cv_file_path,
        filename=f"CV_{candidate.first_name}_{candidate.last_name}{Path(candidate.cv_file_path).suffix}",
        media_type="application/pdf"
    )


def analyze_job_candidate_match_with_llm(cv_text: str, job_data: dict) -> dict:
    """Utilise un LLM pour analyser en profondeur la correspondance entre un CV et un besoin de recrutement"""
    if genai is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="google-generativeai n'est pas installé. Installez-le avec: pip install google-generativeai"
        )
    
    # Récupérer la clé API depuis les variables d'environnement
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("OPENAI_API_KEY")  # Compatibilité avec ancienne variable
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="GEMINI_API_KEY n'est pas configurée dans les variables d'environnement"
        )
    
    # Configurer Gemini
    genai.configure(api_key=api_key)
    # Utiliser gemini-1.5-pro (modèle stable et disponible)
    model = genai.GenerativeModel('gemini-1.5-pro')
    
    # Construire une description structurée du besoin
    job_description = f"""
POSTE: {job_data.get('title', 'N/A')}
Département: {job_data.get('department', 'N/A')}
Type de contrat: {job_data.get('contract_type', 'N/A')}

MISSIONS PRINCIPALES:
{job_data.get('missions_principales', 'Non spécifié')}

MISSIONS SECONDAIRES:
{job_data.get('missions_secondaires', 'Non spécifié')}

PROFIL RECHERCHÉ:
- Niveau de formation: {job_data.get('niveau_formation', 'Non spécifié')}
- Expérience requise: {job_data.get('experience_requise', 'Non spécifié')} ans
- Compétences techniques obligatoires: {', '.join(job_data.get('competences_techniques_obligatoires', []) or [])}
- Compétences techniques souhaitées: {', '.join(job_data.get('competences_techniques_souhaitees', []) or [])}
- Compétences comportementales: {', '.join(job_data.get('competences_comportementales', []) or [])}
- Langues requises: {job_data.get('langues_requises', 'Non spécifié')}
- Certifications requises: {job_data.get('certifications_requises', 'Non spécifié')}

CONTRAINTES:
- Localisation: {job_data.get('localisation', 'Non spécifié')}
- Mobilité: {job_data.get('mobilite_deplacements', 'Non spécifié')}
- Télétravail: {job_data.get('teletravail', 'Non spécifié')}
- Critères éliminatoires: {job_data.get('criteres_eliminatoires', 'Aucun')}
"""
    
    # Prompt pour l'analyse approfondie
    prompt = f"""Tu es un expert en recrutement et en analyse de profils. Analyse en profondeur la correspondance entre le CV d'un candidat et un besoin de recrutement.

BESOIN DE RECRUTEMENT:
{job_description}

CV DU CANDIDAT:
{cv_text}

Effectue une analyse approfondie et détaillée de la correspondance. Retourne un JSON avec la structure suivante:

{{
  "overall_score": nombre entre 0 et 100 (score global de correspondance),
  "overall_assessment": "Évaluation globale détaillée de l'adéquation du candidat au poste (2-3 paragraphes)",
  "technical_skills_analysis": "Analyse approfondie des compétences techniques: correspondances exactes, compétences similaires, lacunes, niveau de maîtrise estimé (3-4 paragraphes)",
  "experience_analysis": "Analyse de l'expérience: pertinence, durée, secteurs, projets similaires, progression de carrière (2-3 paragraphes)",
  "soft_skills_analysis": "Analyse des compétences comportementales si disponibles dans le CV (1-2 paragraphes, ou null si non applicable)",
  "education_analysis": "Analyse de la formation: adéquation du niveau, pertinence du cursus, certifications (1-2 paragraphes, ou null si non applicable)",
  "language_analysis": "Analyse des langues si mentionnées (1 paragraphe, ou null si non applicable)",
  "strengths": ["point fort 1", "point fort 2", ...] (liste de 3-5 points forts),
  "weaknesses": ["point faible 1", "point faible 2", ...] (liste de 2-4 points faibles ou manquants),
  "recommendations": ["recommandation 1", "recommandation 2", ...] (liste de 3-5 recommandations pour l'entretien ou le recrutement),
  "matching_skills": ["compétence 1", "compétence 2", ...] (liste des compétences qui correspondent parfaitement),
  "missing_skills": ["compétence 1", "compétence 2", ...] (liste des compétences manquantes ou à développer),
  "complementary_skills": ["compétence 1", "compétence 2", ...] (liste des compétences complémentaires du candidat non demandées mais utiles),
  "technical_score": nombre entre 0 et 100,
  "experience_score": nombre entre 0 et 100,
  "soft_skills_score": nombre entre 0 et 100 ou null,
  "education_score": nombre entre 0 et 100 ou null,
  "language_score": nombre entre 0 et 100 ou null
}}

Règles importantes:
- Sois précis et détaillé dans tes analyses
- Base-toi sur le contenu réel du CV et du besoin
- Identifie les correspondances exactes ET les compétences transférables
- Mentionne les projets ou expériences pertinents du CV
- Les analyses doivent être en français
- Retourne UNIQUEMENT le JSON, sans texte avant ou après
"""
    
    try:
        # Construire le prompt complet pour Gemini
        full_prompt = f"""Tu es un expert en recrutement spécialisé dans l'analyse approfondie de CV et de besoins de recrutement. Tu retournes uniquement du JSON valide.

{prompt}"""
        
        # Configuration pour la génération
        generation_config = {
            "temperature": 0.3,
            "max_output_tokens": 4000,
        }
        
        response = model.generate_content(
            full_prompt,
            generation_config=generation_config
        )
        
        # Extraire le JSON de la réponse
        response_text = response.text.strip()
        
        # Nettoyer la réponse (enlever les markdown code blocks si présents)
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        # Parser le JSON
        parsed_data = json.loads(response_text)
        
        # Valider les champs obligatoires
        required_fields = ["overall_score", "overall_assessment", "technical_skills_analysis", 
                          "experience_analysis", "strengths", "weaknesses", "recommendations",
                          "matching_skills", "missing_skills", "complementary_skills",
                          "technical_score", "experience_score"]
        for field in required_fields:
            if field not in parsed_data:
                raise ValueError(f"Champ obligatoire manquant: {field}")
        
        return parsed_data
        
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors du parsing JSON de la réponse LLM: {str(e)}. Réponse reçue: {response_text[:500]}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de l'analyse IA: {str(e)}"
        )


@router.post("/{candidate_id}/compare-with-job/{job_id}", response_model=JobCandidateComparisonResponse)
def compare_candidate_with_job(
    candidate_id: UUID,
    job_id: UUID,
    current_user: User = Depends(get_current_active_user),
    session: Session = Depends(get_session)
):
    """
    Analyse IA approfondie de la correspondance entre un candidat et un besoin de recrutement
    
    Cette fonction analyse en profondeur le CV du candidat et le besoin de recrutement
    pour fournir une évaluation détaillée de l'adéquation.
    """
    # Récupérer le candidat
    candidate = session.get(Candidate, candidate_id)
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidat non trouvé"
        )
    
    # Récupérer le besoin
    job = session.get(Job, job_id)
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Besoin de recrutement non trouvé"
        )
    
    # Vérifier que le candidat a un CV
    if not candidate.cv_file_path or not os.path.exists(candidate.cv_file_path):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le candidat doit avoir un CV pour effectuer l'analyse"
        )
    
    # Extraire le texte du CV
    try:
        file_extension = Path(candidate.cv_file_path).suffix.lower()
        if file_extension == ".pdf":
            cv_text = extract_text_from_pdf(candidate.cv_file_path)
        elif file_extension in {".doc", ".docx"}:
            cv_text = extract_text_from_docx(candidate.cv_file_path)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Format de CV non supporté: {file_extension}"
            )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de l'extraction du texte du CV: {str(e)}"
        )
    
    # Préparer les données du besoin pour l'analyse
    job_data = {
        "title": job.title,
        "department": job.department,
        "contract_type": job.contract_type,
        "missions_principales": job.missions_principales,
        "missions_secondaires": job.missions_secondaires,
        "niveau_formation": job.niveau_formation,
        "experience_requise": job.experience_requise,
        "competences_techniques_obligatoires": job.competences_techniques_obligatoires or [],
        "competences_techniques_souhaitees": job.competences_techniques_souhaitees or [],
        "competences_comportementales": job.competences_comportementales or [],
        "langues_requises": job.langues_requises,
        "certifications_requises": job.certifications_requises,
        "localisation": job.localisation,
        "mobilite_deplacements": job.mobilite_deplacements,
        "teletravail": job.teletravail,
        "criteres_eliminatoires": job.criteres_eliminatoires,
    }
    
    # Vérifier si une analyse existe déjà
    existing_comparison = session.exec(
        select(CandidateJobComparison).where(
            CandidateJobComparison.candidate_id == candidate_id,
            CandidateJobComparison.job_id == job_id
        )
    ).first()
    
    # Effectuer l'analyse IA
    analysis_result = analyze_job_candidate_match_with_llm(cv_text, job_data)
    
    # Sauvegarder ou mettre à jour l'analyse
    analysis_json = json.dumps(analysis_result, ensure_ascii=False)
    
    try:
        if existing_comparison:
            # Mettre à jour l'analyse existante
            existing_comparison.analysis_data = analysis_json
            existing_comparison.updated_at = datetime.utcnow()
            existing_comparison.created_by = current_user.id
            session.add(existing_comparison)
            session.commit()
            session.refresh(existing_comparison)
        else:
            # Créer une nouvelle analyse
            new_comparison = CandidateJobComparison(
                candidate_id=candidate_id,
                job_id=job_id,
                created_by=current_user.id,
                analysis_data=analysis_json
            )
            session.add(new_comparison)
            session.commit()
            session.refresh(new_comparison)
    except Exception as save_error:
        # Logger l'erreur et lever une exception pour que l'utilisateur soit informé
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"❌ ERREUR CRITIQUE lors de la sauvegarde de l'analyse IA pour candidat {candidate_id} et job {job_id}: {str(save_error)}", exc_info=True)
        # Rollback et lever une exception pour informer l'utilisateur
        try:
            session.rollback()
        except:
            pass
        # Lever une exception pour que l'utilisateur sache que la sauvegarde a échoué
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"L'analyse IA a été générée mais n'a pas pu être sauvegardée. Erreur: {str(save_error)}. Veuillez vérifier que la table 'candidate_job_comparisons' existe dans la base de données."
        )
    
    # Convertir en réponse
    return JobCandidateComparisonResponse(**analysis_result)


@router.get("/{candidate_id}/compare-with-job/{job_id}", response_model=JobCandidateComparisonResponse)
def get_saved_comparison(
    candidate_id: UUID,
    job_id: UUID,
    current_user: User = Depends(get_current_active_user),
    session: Session = Depends(get_session)
):
    """
    Récupère une analyse IA sauvegardée pour un candidat et un besoin
    """
    # Vérifier que le candidat et le besoin existent
    candidate = session.get(Candidate, candidate_id)
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidat non trouvé"
        )
    
    job = session.get(Job, job_id)
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Besoin de recrutement non trouvé"
        )
    
    # Récupérer l'analyse sauvegardée
    comparison = session.exec(
        select(CandidateJobComparison).where(
            CandidateJobComparison.candidate_id == candidate_id,
            CandidateJobComparison.job_id == job_id
        )
    ).first()
    
    if not comparison:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Aucune analyse sauvegardée trouvée. Veuillez effectuer une analyse d'abord."
        )
    
    # Désérialiser les données JSON
    try:
        analysis_data = json.loads(comparison.analysis_data)
        return JobCandidateComparisonResponse(**analysis_data)
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de la lecture de l'analyse sauvegardée: {str(e)}"
        )
